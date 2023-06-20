from docx import Document
import os

FOLDER = "1eKOOK"


def points_to_you(expected, result):
    ex = expected.split("-")
    if ex[0] == result[0] and ex[1] == result[1]:
        return 15
    elif result[0] > result[1] and ex[0] > ex[1] or\
        result[0] < result[1] and ex[0] < ex[1] or\
        result[0] == result[1] and ex[0] == ex[1]:
        return 10
    elif ex[0] == result[0] and ex[1] != result[1] or\
        ex[0] != result[0] and ex[1] == result[1]:
        return 5
    return 0


def gfa(naam):
    with open('groepsfase.txt', 'r') as fp:
        lines = fp.readlines()
        for row in lines:
            if row.find(naam) != -1:
                points = row.split()[-1]
                return int(points)
    print("NIKS GEVONDEN G FIKS HET!! PROBLEEM: " + naam)


def amount_of_points(filename):
    wordDoc = Document(FOLDER + '/' + filename)
    iterator = 0
    j = 0
    points = 0
    for table in wordDoc.tables:
        for row in table.rows:
            # print(row.cells[0].text)
            if iterator == 0:
                hole_block = row.cells[0].text
                naam = hole_block.replace("Naam en mailadres: ", '')
                naam = naam.rstrip()
                points = gfa(naam)
            if iterator > 1 and iterator < 10:
                if j < len(result):
                    points += points_to_you(row.cells[2].text, result[j])
                else:
                    break
                j += 1
            iterator += 1
    return naam, points

def makefile(tussenstand):
    doc = Document()
    doc.add_heading('Tussenstand WK Poule', 0)
    tussenstand.sort(key=lambda a: (a[1], a[0]), reverse=True)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ranking'
    hdr_cells[1].text = 'Naam'
    hdr_cells[2].text = 'Punten'
    it = 1
    for naam, punten in tussenstand:
        row_cells = table.add_row().cells
        row_cells[0].text = str(it)
        row_cells[1].text = naam
        row_cells[2].text = str(punten)
        it += 1
    doc.add_page_break()
    doc.save('KO.docx')

def groepsfasefile(ranking):
    with open('KO.txt', 'w') as f:
        for item in ranking:
            f.write(item[0] + " " + str(item[1]) + "\n")

f = open("uitslagen","r")
lines = f.readlines()
result = []
for item in lines:
    temp = item.split(" ")
    if len(temp) > 2:
        result.append([temp[2], temp[3].strip()])


lijst = os.listdir(FOLDER)
ranking = []
for file in lijst: 
    ranking.append(amount_of_points(file))
groepsfasefile(ranking)
print(ranking)
for item in ranking:
    print(item[0] + " " + str(item[1]))
# makefile(ranking)

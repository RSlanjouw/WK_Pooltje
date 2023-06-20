from docx import Document

wordDoc = Document('Bonuspunten overzicht.docx')
iterator = 0

def gfa(naam):
    with open('KO2.txt', 'r') as fp:
        lines = fp.readlines()
        for row in lines:
            if row.find(naam) != -1:
                points = row.split()[-1]
                return int(points)
    print("NIKS GEVONDEN G FIKS HET!! PROBLEEM: " + naam)

def bonus_gele(aantal, ag = 227):
    miss = abs(aantal - ag)
    if (100 - miss * 3) > 0:
        return 100 - miss * 3
    return 0

def bonus_rood(aantal, ag = 4):
    miss = abs(aantal - ag)
    if (50 - miss * 3) > 0:
        return 50 - miss * 3
    return 0

def bonus_goals(aantal, ag = 172):
    miss = abs(aantal - ag)
    if (100 - miss * 3) > 0:
        return 100 - miss * 3
    return 0



def makefile(tussenstand):
    doc = Document()
    doc.add_heading('Einduitslag WK Poule', 0)
    tussenstand.sort(key=lambda a: (a[1], a[0]), reverse=True)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ranking'
    hdr_cells[1].text = 'Naam'
    hdr_cells[2].text = 'Punten'
    it = 1
    for naam, punten in tussenstand:
        row_cells = table.add_row().cells
        row_cells[0].text = str(it)
        row_cells[1].text = naam
        row_cells[2].text = str(punten)
        it += 1
    doc.add_page_break()
    doc.save('Einduitslag.docx')
 

ranking = []
for table in wordDoc.tables:
    for row in table.rows:
        if iterator > 0 and iterator < 54:
            print(row.cells[0].text)
            naam = row.cells[0].text
            points = gfa(row.cells[0].text.split()[0])
            points += bonus_gele(int(row.cells[1].text))
            points += bonus_rood(int(row.cells[2].text))
            points += bonus_goals(int(row.cells[3].text))
            points += int(row.cells[5].text)
            points += int(row.cells[4].text)
            ranking.append([naam, points])
        elif iterator > 53:
            # ranking.append([naam, points])
            break
        # print(row.cells[0].text)
        # for cell in row.cells:
            # print(cell.text)
        print("--------------------------------")
        iterator += 1
makefile(ranking)
print(ranking)
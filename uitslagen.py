from docx import Document
import os

FOLDER = "2eKO"

def halve_finale(expected, result, Landen):
    ex = expected.split("-")
    print(ex)
    if ex[0] == result[0] and ex[1] == result[1] and (int(result[2]) in Landen) and (int(result[3]) in Landen):
        return 20
    elif (result[0] > result[1] and ex[0] > ex[1] or
        result[0] < result[1] and ex[0] < ex[1] or
        result[0] == result[1] and ex[0] == ex[1]) and (int(result[2]) in Landen) and (int(result[3]) in Landen):
        return 15
    elif ex[0] == result[0] and ex[1] == result[1] and ((int(result[2]) in Landen) or (int(result[3]) in Landen)):
        return 10
    elif (result[0] > result[1] and ex[0] > ex[1] or
        result[0] < result[1] and ex[0] < ex[1] or
        result[0] == result[1] and ex[0] == ex[1]) and ((int(result[2]) in Landen) or (int(result[3]) in Landen)):
        return 5
    return 0

def points_to_you(expected, result):
    ex = expected.split("-")
    if ex[0] == result[0] and ex[1] == result[1]:
        return 20
    elif result[0] > result[1] and ex[0] > ex[1] or\
        result[0] < result[1] and ex[0] < ex[1] or\
        result[0] == result[1] and ex[0] == ex[1]:
        return 15
    elif ex[0] == result[0] and ex[1] != result[1] or\
        ex[0] != result[0] and ex[1] == result[1]:
        return 5
    return 0


def gfa(naam):
    with open('KO.txt', 'r') as fp:
        lines = fp.readlines()
        for row in lines:
            if row.find(naam) != -1:
                points = row.split()[-1]
                return int(points)
    print("NIKS GEVONDEN G FIKS HET!! PROBLEEM: " + naam)


def deze_eruit(uitslag, row, j):
    ex = uitslag.split("-")
    if ex[0] < ex[1]:
        return 1 + j * 2
    elif ex[1] < ex[0]:
        return 2 + j * 2
    elif ex[0] == ex[1]:
        # print("De voorspelling is:" + row.cells[4].text)
        if row.cells[4].text == "Nederland":
            return 2
        if row.cells[4].text == "Frankrijk":
            return 5
        if row.cells[4].text == "Portugal":
            return 7
        return int(input("De voorspelling is: " + row.cells[4].text + " J is:" + str(j))) 


def amount_of_points(filename):
    wordDoc = Document(FOLDER + '/' + filename)
    iterator = 0
    j = 0
    points = 0
    Landen = [1,2,3,4,5,6,7,8]
    for table in wordDoc.tables:
        for row in table.rows:
            # print(row.cells[0].text)
            if iterator == 0:
                hole_block = row.cells[0].text
                naam = hole_block.replace("NAAM: ", '')
                naam = naam.strip()
                points = gfa(naam)
                print(naam)
                # iterator += 2
            if iterator > 1 and iterator % 2 == 0:
                # print(iter)
                if j < 4:
                    print(row.cells[3].text)
                    points += points_to_you(row.cells[3].text, result[j])
                    Landen.remove(deze_eruit(row.cells[3].text, row, j))
                elif j < 6:
                    points += halve_finale(row.cells[3].text, result[j], Landen)
                else:
                    break
                j += 1
                iterator += 2            
            iterator += 1
    return naam, points

def makefile(tussenstand):
    doc = Document()
    doc.add_heading('Tussenstand WK Poule', 0)
    tussenstand.sort(key=lambda a: (a[1], a[0]), reverse=True)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ranking'
    hdr_cells[1].text = 'Naam'
    hdr_cells[2].text = 'Punten'
    it = 1
    for naam, punten in tussenstand:
        row_cells = table.add_row().cells
        row_cells[0].text = str(it)
        row_cells[1].text = naam
        row_cells[2].text = str(punten)
        it += 1
    doc.add_page_break()
    doc.save('KO.docx')

def groepsfasefile(ranking):
    with open('KO2.txt', 'w') as f:
        for item in ranking:
            f.write(item[0] + " " + str(item[1]) + "\n")

f = open("uitslagen","r")
lines = f.readlines()
result = []
for item in lines:
    temp = item.split(" ")
    if len(temp) > 2:
        result.append([temp[2], temp[3].strip(), temp[0], temp[1]])

lijst = os.listdir(FOLDER)
ranking = []
for file in lijst: 
    ranking.append(amount_of_points(file))
groepsfasefile(ranking)
makefile(ranking)
